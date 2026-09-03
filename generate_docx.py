import docx
from docx.shared import Inches, Pt
import os

doc = docx.Document()
doc.add_heading('Assignment Solution: UniCore Resource Orchestration', 0)

doc.add_heading('1. Problem Understanding and Formulation', level=1)
doc.add_paragraph(
    "Scenario: UniCore is a smart university computing platform supporting 500-800 users. "
    "It provides interactive and background services. The challenge is to efficiently orchestrate "
    "CPU, memory, files, and disk I/O while ensuring synchronization, avoiding deadlocks, and preventing thrashing."
)
doc.add_paragraph(
    "Assumptions:\n"
    "- Workload: 4 interactive processes, 4 background processes.\n"
    "- Resources: 4 types (CPU, RAM, DB Locks, Network Ports).\n"
    "- Memory: 16 GB RAM, 4 KB pages.\n"
    "- Disk: 200 cylinders, 10 pending requests."
)

doc.add_heading('2. Process Management, Synchronization & Deadlock', level=1)
doc.add_paragraph("Synchronization Pseudocode (Readers-Writers):")
doc.add_paragraph(
    "semaphore rw_mutex = 1;\n"
    "semaphore mutex = 1;\n"
    "int read_count = 0;\n\n"
    "Writer:\n"
    "wait(rw_mutex);\n"
    "// write data\n"
    "signal(rw_mutex);\n\n"
    "Reader:\n"
    "wait(mutex);\n"
    "read_count++;\n"
    "if (read_count == 1) wait(rw_mutex);\n"
    "signal(mutex);\n"
    "// read data\n"
    "wait(mutex);\n"
    "read_count--;\n"
    "if (read_count == 0) signal(rw_mutex);\n"
    "signal(mutex);"
)
doc.add_paragraph("Banker's Algorithm Output:")
if os.path.exists("bankers_output.png"):
    doc.add_picture("bankers_output.png", width=Inches(6.0))
else:
    doc.add_paragraph("[Banker's Algorithm Screenshot Here]")

doc.add_heading('3. Memory Management', level=1)
doc.add_paragraph(
    "Frames available: 16 GB / 4 KB = (16 * 1024^3) / 4096 = 4,194,304 frames.\n"
    "Pages required for 64 MB logical address space: 64 MB / 4 KB = (64 * 1024^2) / 4096 = 16,384 pages."
)
doc.add_paragraph("Page Replacement Output:")
if os.path.exists("memory_output.png"):
    doc.add_picture("memory_output.png", width=Inches(6.0))
else:
    doc.add_paragraph("[Memory Management Screenshot Here]")
    
doc.add_heading('4. File Systems & Disk Scheduling', level=1)
doc.add_paragraph("File Allocation: Contiguous (fast sequential, high fragmentation), Linked (no external fragmentation, slow random), Indexed (good random, overhead). Recommended: Indexed Allocation.")
doc.add_paragraph("Disk Scheduling Output:")
if os.path.exists("disk_output.png"):
    doc.add_picture("disk_output.png", width=Inches(6.0))
else:
    doc.add_paragraph("[Disk Scheduling Screenshot Here]")
    
doc.add_heading('5. CPU Scheduling', level=1)
doc.add_paragraph("CPU Scheduling Output (FCFS vs Round Robin):")
if os.path.exists("cpu_output.png"):
    doc.add_picture("cpu_output.png", width=Inches(6.0))
else:
    doc.add_paragraph("[CPU Scheduling Screenshot Here]")

doc.add_heading('6. Conclusion & Reflection', level=1)
doc.add_paragraph("By integrating these OS concepts, the UniCore system achieves fair CPU scheduling through Round Robin, avoids deadlocks via Banker's algorithm, manages memory effectively, and uses SSTF/Indexed allocation for fast I/O.")

doc.save('Solution.docx')
print("Solution.docx generated successfully.")
